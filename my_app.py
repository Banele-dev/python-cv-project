""""
# Automatic YouTube Video Downloader 
import pytube 

link = input('Youtube Video URL')
video_download = pytube.YouTube(link)

# Download the video to my computer.
video_download.streams.first().download()

print('Video Downloaded', link) 

"""
"""
# Checking for the Internet Speed 
import speedtest as st 

# Set Best Server 
server = st.Speedtest()
server.get_best_server()

# Test Download Speed
download = server.download()
download = download / 1000000
print(f"Download Speed: {download} MB/s")

# Test Upload Speed
upload = server.upload()
upload = upload / 1000000
print(f"Upload Speed: {upload} Mb/s")

# Test ping || This checks if the computer can communicate with another computer/server over a network.
ping = server.results.ping
print(f"Ping Speed: {ping}")

"""
"""
# Python script for web scraping to extract data from a website
import requests 
from bs4 import BeautifulSoup

session = requests.session()
req = session.get('https://stackoverflow.com/questions/10807081/script-to-extract-data-from-web-page')
doc = BeautifulSoup(req.content, 'html.parser') # It fetches the content of the webpage and uses BeautifulSoup to parse the HTML.
print(doc.findAll('a', {"class" : "gp-share"}))

print(doc)

"""

from docx import Document
# from docx.shared import Inches
import pyttsx3

def speak(text):
    pyttsx3.speak(text)

document = Document()

# document.add_picture('OIP.jpg', width=Inches(2.0))

# name phone number email address detail
name = input('What is your name? ')
speak('Hello ' + name + ' how are you today?')

speak('What is your phone number?')
phone_number = input('What is your phone number? ')
email = input('What is your email address? ')

document.add_paragraph(
    name + ' | ' + phone_number + ' | ' + email)

# About me
document.add_heading('About me')
document.add_paragraph(
    input('Tell me about yourself? ')
)

# Work expereince
document.add_heading('Work Experience')
p = document.add_paragraph()

company = input('Enter company name ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + ' ' + to_date + '\n').italic = True

experience_details = input(
    'Describe your expereince at ' + company)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        'Do you have more job experiences? Yes or No ')
    if has_more_experiences.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter company name ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + ' ' + to_date + '\n').italic = True

        experience_details = input(
            'Describe your expereince at ' + company + ' ' )
        p.add_run(experience_details) 
    else:
        break
    
# Skills 
document.add_heading('Skills')
skill = input('Enter your skill')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_Skills = input('Do you have more skills? Yes or No ')
    if has_more_Skills.lower() == 'yes':
        skill = input('Enter your skill')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# Footer 
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV genereted using Youtube"

document.save('cv.docx')



    
